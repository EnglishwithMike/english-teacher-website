"""Private Supabase Storage helpers for LearningXY."""

import html
import io
import mimetypes
import os
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document

from flask import Response
from supabase import create_client
from dotenv import load_dotenv

load_dotenv()


PROOF_BUCKET = "teacher-proofs"
PROFILE_BUCKET = "teacher-profile-images"


def _client():
    return create_client(
        os.environ["SUPABASE_URL"],
        os.environ["SUPABASE_SECRET_KEY"],
    )


def upload_file(bucket, filename, file_object, content_type=None):
    file_object.seek(0)
    data = file_object.read()
    content_type = (
        content_type
        or getattr(file_object, "mimetype", None)
        or mimetypes.guess_type(filename)[0]
        or "application/octet-stream"
    )
    _client().storage.from_(bucket).upload(
        filename,
        data,
        {"content-type": content_type, "upsert": "true"},
    )
    file_object.seek(0)


def download_file(bucket, filename):
    return _client().storage.from_(bucket).download(filename)


def remove_file(bucket, filename):
    if not filename:
        return
    try:
        _client().storage.from_(bucket).remove([filename])
    except Exception as error:
        # A missing old object should not prevent account/application deletion.
        print("SUPABASE STORAGE REMOVE ERROR:", repr(error))


def inline_response(bucket, filename, download_name=None):
    data = download_file(bucket, filename)
    mimetype = mimetypes.guess_type(download_name or filename)[0]
    response = Response(data, mimetype=mimetype or "application/octet-stream")
    response.headers["Content-Disposition"] = (
        f'inline; filename="{download_name or filename}"'
    )
    response.headers["X-Content-Type-Options"] = "nosniff"
    return response


def _document_html_response(title, body_html):
    page = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{html.escape(title)}</title>
    <style>
        body {{
            margin: 0;
            padding: 28px;
            background: #eef2fb;
            color: #17213b;
            font-family: Arial, sans-serif;
            line-height: 1.55;
        }}
        .document {{
            max-width: 850px;
            min-height: 900px;
            margin: 0 auto;
            padding: 45px;
            box-sizing: border-box;
            background: white;
            border: 1px solid #ccd4e7;
            box-shadow: 0 8px 24px rgba(0, 0, 0, .12);
        }}
        h1, h2, h3, p {{ overflow-wrap: anywhere; }}
        table {{
            width: 100%;
            margin: 16px 0;
            border-collapse: collapse;
        }}
        td {{
            padding: 8px;
            border: 1px solid #aeb9d3;
            vertical-align: top;
        }}
        .notice {{
            padding: 14px;
            border-left: 5px solid #243b72;
            background: #f2f5fc;
        }}
    </style>
</head>
<body><main class="document">{body_html}</main></body>
</html>"""
    response = Response(page, mimetype="text/html")
    response.headers["Content-Disposition"] = "inline"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["Content-Security-Policy"] = (
        "default-src 'none'; style-src 'unsafe-inline'; img-src data:"
    )
    return response


def _docx_preview(data, filename):
    document = Document(io.BytesIO(data))
    sections = []

    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue

        escaped = html.escape(value)
        style_name = (paragraph.style.name or "").lower()
        if "title" in style_name:
            sections.append(f"<h1>{escaped}</h1>")
        elif "heading 1" in style_name:
            sections.append(f"<h2>{escaped}</h2>")
        elif "heading" in style_name:
            sections.append(f"<h3>{escaped}</h3>")
        elif "list" in style_name:
            sections.append(f"<p>• {escaped}</p>")
        else:
            sections.append(f"<p>{escaped}</p>")

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = "".join(
                f"<td>{html.escape(cell.text.strip())}</td>"
                for cell in row.cells
            )
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            sections.append("<table>" + "".join(rows) + "</table>")

    if not sections:
        sections.append(
            '<p class="notice">The document contains no readable text.</p>'
        )

    return _document_html_response(filename, "".join(sections))


def _odt_preview(data, filename):
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        content = archive.read("content.xml")

    root = ElementTree.fromstring(content)
    sections = []

    for element in root.iter():
        local_name = element.tag.rsplit("}", 1)[-1]
        if local_name not in {"p", "h"}:
            continue

        value = "".join(element.itertext()).strip()
        if not value:
            continue

        escaped = html.escape(value)
        if local_name == "h":
            sections.append(f"<h2>{escaped}</h2>")
        else:
            sections.append(f"<p>{escaped}</p>")

    if not sections:
        sections.append(
            '<p class="notice">The document contains no readable text.</p>'
        )

    return _document_html_response(filename, "".join(sections))


def proof_preview_response(filename):
    extension = filename.rsplit(".", 1)[-1].lower()

    if extension in {"pdf", "png", "jpg", "jpeg"}:
        return inline_response(PROOF_BUCKET, filename)

    data = download_file(PROOF_BUCKET, filename)

    try:
        if extension == "docx":
            return _docx_preview(data, filename)

        if extension == "odt":
            return _odt_preview(data, filename)
    except (ValueError, KeyError, zipfile.BadZipFile):
        return _document_html_response(
            filename,
            '<p class="notice">The document is damaged or could not be read.</p>',
        )

    if extension == "doc":
        return _document_html_response(
            filename,
            '<p class="notice">This older Word document cannot be displayed '
            'safely in the browser. Ask the applicant to upload it as PDF, '
            'DOCX or ODT.</p>',
        )

    return "This file type cannot be previewed.", 400


def cache_profile_image(filename, application_root):
    if not filename:
        return

    directory = Path(application_root) / "static" / "teacher_profiles"
    directory.mkdir(parents=True, exist_ok=True)
    destination = directory / Path(filename).name

    if destination.is_file():
        return

    try:
        destination.write_bytes(download_file(PROFILE_BUCKET, filename))
    except Exception as error:
        print("SUPABASE PROFILE CACHE ERROR:", repr(error))
